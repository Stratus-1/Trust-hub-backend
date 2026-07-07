from fastapi import APIRouter, HTTPException, Request
from docx import Document
from io import BytesIO
import os
from fastapi.responses import StreamingResponse

router = APIRouter()

@router.post("/generate-resolution")
async def generate_resolution(request: Request):
    try:
        trust = await request.json()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Invalid JSON: {e}")

    # Validate required fields
    required_fields = ["email", "trust_name", "trust_number", "full_name"]
    missing_fields = [field for field in required_fields if not trust.get(field)]
    if missing_fields:
        raise HTTPException(status_code=400, detail=f"Missing required fields: {', '.join(missing_fields)}")

    # Default trustee names to empty strings
    for key in ['trustee1_name', 'trustee2_name', 'trustee3_name', 'trustee4_name']:
        trust[key] = trust.get(key, "")

    # Choose template
    trustee_count = sum(1 for k in ['trustee1_name', 'trustee2_name', 'trustee3_name', 'trustee4_name'] if trust.get(k))
    template_path = f"templates/HK_Trust_Trustee_Appointment_Resolution_{trustee_count}_Trustees.docx"
    if not os.path.exists(template_path):
        raise HTTPException(status_code=500, detail="Missing resolution template.")

    # Helper function to replace placeholders in paragraphs
    def replace_placeholders_in_paragraphs(paragraphs, mapping):
        for paragraph in paragraphs:
            for key, val in mapping.items():
                if f"{{{{{key}}}}}" in paragraph.text:
                    for run in paragraph.runs:
                        if f"{{{{{key}}}}}" in run.text:
                            run.text = run.text.replace(f"{{{{{key}}}}}", str(val))

    # Load template and fill fields
    try:
        doc = Document(template_path)
        replace_placeholders_in_paragraphs(doc.paragraphs, trust)

        for section in doc.sections:
            replace_placeholders_in_paragraphs(section.header.paragraphs, trust)
            replace_placeholders_in_paragraphs(section.footer.paragraphs, trust)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Document generation failed: {e}")

    # Save to memory
    try:
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error saving document: {e}")

    filename = f"Trustee_Resolution_{trust['trust_number'].replace('/', '_')}.docx"
    return StreamingResponse(
        content=doc_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename=\"{filename}\"; filename*=UTF-8''{filename}"
        }
    )
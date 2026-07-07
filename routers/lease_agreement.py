from fastapi import APIRouter, Request
from fastapi.responses import FileResponse
from docx import Document
from tempfile import NamedTemporaryFile
from datetime import date
import shutil

router = APIRouter()

@router.post("/generate-lease-agreement")
async def generate_lease_agreement(data: dict):
    template_path = "templates/HK_Foreign_Trust_Lease_Agreement.docx"
    doc = Document(template_path)

    today = date.today().strftime("%d %B %Y")

    merge_fields = {
        "trust_name": data.get("trust_name", ""),
        "trust_number": data.get("trust_number", ""),
        "owner_name": data.get("owner_name", ""),
        "owner_id": data.get("owner_id", ""),
        "signer_name": data.get("signer_name", ""),
        "Property_Address": data.get("Property_Address", ""),
        "witness_name": data.get("witness_name", ""),
        "witness_id": data.get("witness_id", ""),
        "establishment_date_1": data.get("establishment_date_1", today),
        "establishment_date_2": data.get("establishment_date_2", today)
    }

    def merge(doc_obj):
        def replace_text_in_runs(paragraph):
            full_text = "".join(run.text for run in paragraph.runs)
            replaced = False
            for key, value in merge_fields.items():
                if f"{{{{{key}}}}}" in full_text:
                    full_text = full_text.replace(f"{{{{{key}}}}}", str(value))
                    replaced = True
            if replaced:
                # Assign all to first run and clear others
                if paragraph.runs:
                    paragraph.runs[0].text = full_text
                    for run in paragraph.runs[1:]:
                        run.text = ""

        for p in doc_obj.paragraphs:
            replace_text_in_runs(p)

        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_text_in_runs(p)

        for section in doc_obj.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                for p in header.paragraphs:
                    replace_text_in_runs(p)
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                for p in footer.paragraphs:
                    replace_text_in_runs(p)

    merge(doc)

    with NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        return FileResponse(path=tmp.name, filename="HK_Lease_Agreement.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")